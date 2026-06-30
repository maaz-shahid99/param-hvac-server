#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Focused editable .docx: 3-platform comparison on Scalability / Technical
Feasibility / Costing, grounded in ~20 GB data from SAP, long-term optimization."""
import struct, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DIAG = r"c:\Users\maazs\Documents\Projects\HVAC_v1.1\_diagrams"
OUT = r"c:\Users\maazs\Documents\Projects\HVAC_v1.1\BI_Comparison_Dollar_Industries.docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
TEAL = RGBColor(0x12, 0x6E, 0x82)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "1F3A5F"
GROUP_FILL = "126E82"
VERDICT_FILL = "E7EEF3"
ZEBRA = "F2F6FA"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(10.5)
for lvl, sz, col in [("Title", 24, NAVY), ("Heading 1", 14, NAVY), ("Heading 2", 12, TEAL)]:
    st = doc.styles[lvl]; st.font.name = "Calibri"; st.font.size = Pt(sz)
    st.font.color.rgb = col; st.font.bold = True


def png_size(path):
    with open(path, "rb") as f:
        head = f.read(26)
    return struct.unpack(">II", head[16:24])


def add_image(path, max_w=6.4, max_h=7.5, caption=None):
    w, h = png_size(path); aspect = h / w; width = max_w
    if width * aspect > max_h:
        width = max_h / aspect
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    if caption:
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY


def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def set_cell(cell, text, bold=False, color=None, size=9.5, italic=False, align=None):
    cell.text = ""; p = cell.paragraphs[0]
    if align: p.alignment = align
    run = p.add_run(str(text)); run.bold = bold; run.italic = italic; run.font.size = Pt(size)
    if color: run.font.color.rgb = color


def para(text=None, bold_lead=None, size=10.5):
    p = doc.add_paragraph()
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True; r.font.size = Pt(size)
    if text:
        r = p.add_run(text); r.font.size = Pt(size)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
    p.add_run(text); return p

# ----------------- TITLE -----------------
doc.add_paragraph("Analytics Platform Comparison", style="Title")
s = doc.add_paragraph()
r = s.add_run("Prepared for Dollar Industries  |  Tableau  vs.  AIBI (our platform)  vs.  GenAI Conversational Assistant")
r.italic = True; r.font.size = Pt(11); r.font.color.rgb = TEAL

# ----------------- SCENARIO BOX -----------------
sc = doc.add_table(rows=1, cols=3); sc.alignment = WD_TABLE_ALIGNMENT.CENTER; sc.style = "Table Grid"
labels = [("Data Volume", "~ 20 GB"), ("Source System", "SAP"), ("Optimization Horizon", "Long-term (3-5 yrs)")]
for j, (k, v) in enumerate(labels):
    c = sc.rows[0].cells[j]; c.text = ""
    p1 = c.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p1.add_run(k + "\n"); rr.font.size = Pt(8.5); rr.font.color.rgb = WHITE; rr.bold = True
    rr2 = p1.add_run(v); rr2.font.size = Pt(12); rr2.font.color.rgb = WHITE; rr2.bold = True
    shade(c, HEADER_FILL)
doc.add_paragraph()

para("At ~20 GB, all three platforms handle the volume comfortably, so raw scalability is not the "
     "differentiator. The long-term differences are driven by SAP connectivity effort and how cost "
     "behaves as the user base grows.")

# ----------------- MAIN COMPARISON TABLE -----------------
doc.add_heading("Comparison Matrix", level=1)
cols = ["Dimension", "Tableau", "AIBI (ours)", "GenAI Assistant"]
rows = [
    ("group", "SCALABILITY"),
    ("row", ["Handling 20 GB today", "Comfortable - Hyper extract or live connect to SAP HANA/BW",
             "Comfortable - 20 GB is trivial for PostgreSQL on a modest instance",
             "Comfortable - warehouse handles 20 GB easily"]),
    ("row", ["Growth headroom (100 GB+)", "Scales, but extract refresh + server sizing cost rises",
             "Good - vertical/horizontal Postgres scaling", "Highest - cloud-native auto-scale, parallel ETL"]),
    ("row", ["Concurrency / many users", "Proven, but each concurrent user = a licensed seat",
             "App tier scales, no seat cap", "Serverless auto-scale, no seat cap"]),
    ("verdict", "Verdict: Fine at 20 GB across the board; for a growing user base, AIBI/GenAI scale more economically than Tableau."),

    ("group", "TECHNICAL FEASIBILITY"),
    ("row", ["SAP connectivity", "Native certified connectors (HANA, BW, NetWeaver) - fastest path",
             "No native SAP connector - needs ETL (SAP to SLT/Data Services/ADF to PostgreSQL)",
             "Proven SAP ingestion (REST adapters / Self-Hosted IR to warehouse), ETL must be built"]),
    ("row", ["Setup effort", "Low (connect + model)", "Medium (build SAP to Postgres pipeline once)",
             "High (full ETL + Bronze/Gold + agents)"]),
    ("row", ["Capability depth", "Dashboards (NL limited)", "Dashboards + conversational",
             "Agentic chat + documents + web"]),
    ("row", ["Maintenance", "Vendor-managed", "We maintain - full control", "Heavier pipeline ops"]),
    ("verdict", "Verdict: Tableau is easiest to connect to SAP; AIBI/GenAI need an ingestion layer but unlock conversational analytics and full customisation."),

    ("group", "COSTING (long-term optimization)"),
    ("row", ["Licensing model", "Recurring per-seat (USD)", "No per-seat - our IP", "No per-seat - custom"]),
    ("row", ["Infrastructure at 20 GB", "Server/Cloud + extract storage", "Small Postgres - low, commodity cost",
             "Warehouse + storage + AI Search - higher fixed"]),
    ("row", ["Variable cost", "None at core", "LLM tokens (controllable - small models + cache)",
             "LLM tokens + RAG indexing (scales with usage)"]),
    ("row", ["Cost curve as you grow", "Rises linearly with users", "Flattens - front-load build, low marginal cost",
             "Moderate fixed, usage-variable"]),
    ("row", ["3-5 yr TCO (this scenario)", "Medium-High", "Low-Medium (best)", "High"]),
    ("verdict", "Verdict: For long-term cost optimization, AIBI is most optimised; Tableau's per-seat model is the main long-term cost risk."),
]

t = doc.add_table(rows=0, cols=4); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
# header
hc = t.add_row().cells
for j, h in enumerate(cols):
    set_cell(hc[j], h, bold=True, color=WHITE, size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER if j else None)
    shade(hc[j], HEADER_FILL)

zebra_i = 0
for kind, payload in rows:
    cells = t.add_row().cells
    if kind == "group":
        m = cells[0]
        for c in cells[1:]:
            m = m.merge(c)
        set_cell(m, payload, bold=True, color=WHITE, size=10)
        shade(m, GROUP_FILL); zebra_i = 0
    elif kind == "verdict":
        m = cells[0]
        for c in cells[1:]:
            m = m.merge(c)
        set_cell(m, payload, italic=True, size=9, color=NAVY)
        shade(m, VERDICT_FILL)
    else:
        for j, val in enumerate(payload):
            set_cell(cells[j], val, bold=(j == 0), size=9)
            if zebra_i % 2 == 1:
                shade(cells[j], ZEBRA)
        zebra_i += 1

for row in t.rows:
    row.cells[0].width = Inches(1.6)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Bottom line: ").bold = True
p.add_run("Tableau is the quickest to connect to SAP, but its per-seat licensing is the main long-term "
          "cost risk as adoption widens. AIBI is the most cost-optimized over time - it needs a one-time "
          "SAP to PostgreSQL ingestion pipeline, then carries no per-seat tax and scales cheaply well "
          "beyond 20 GB.")

# ----------------- NOTES -----------------
doc.add_heading("Notes & Assumptions", level=1)
bullet("At ~20 GB the data layer is inexpensive for all three; cost is driven by licensing model and "
       "(for B/C) LLM token usage, not data volume.")
bullet("Tableau seat pricing is indicative list pricing (Creator ~$75 / Explorer ~$42 / Viewer ~$15 per "
       "user/month, billed annually) and is FX/negotiation sensitive.")
bullet("AIBI and the GenAI Assistant both require a one-time SAP ingestion pipeline (SLT / SAP Data "
       "Services / ADF) since neither connects to SAP natively like Tableau.")
bullet("Candidate C is a restructured reference architecture from another sector; re-scoped here, with "
       "the original client unnamed.")

# ----------------- APPENDIX: ARCHITECTURES -----------------
doc.add_heading("Appendix - Reference Architectures", level=1)
doc.add_heading("A. Tableau", level=2)
add_image(f"{DIAG}\\arch1.png")
doc.add_heading("B. AIBI (our platform)", level=2)
add_image(f"{DIAG}\\arch2.png")
doc.add_heading("C. GenAI Conversational Assistant", level=2)
add_image(f"{DIAG}\\arch3.png")

# ----------------- SAVE (handle file lock) -----------------
try:
    doc.save(OUT); print("SAVED:", OUT)
except PermissionError:
    alt = OUT.replace(".docx", "_v2.docx")
    doc.save(alt); print("ORIGINAL LOCKED - SAVED:", alt)
