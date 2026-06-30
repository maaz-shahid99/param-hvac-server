#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Build an editable .docx pitch deliverable for Dollar Industries."""
import struct
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DIAG = r"c:\Users\maazs\Documents\Projects\HVAC_v1.1\_diagrams"
OUT = r"c:\Users\maazs\Documents\Projects\HVAC_v1.1\BI_Comparison_Dollar_Industries.docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
TEAL = RGBColor(0x12, 0x6E, 0x82)
ACCENT_HEX = "1F3A5F"
HEADER_FILL = "1F3A5F"
ZEBRA = "EEF3F8"

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

for lvl, sz, col in [("Title", 26, NAVY), ("Heading 1", 15, NAVY), ("Heading 2", 12.5, TEAL)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = col
    st.font.bold = True


def png_size(path):
    with open(path, "rb") as f:
        head = f.read(26)
    w, h = struct.unpack(">II", head[16:24])
    return w, h


def add_image(path, max_w=6.4, max_h=8.0, caption=None):
    w, h = png_size(path)
    aspect = h / w
    width = max_w
    if width * aspect > max_h:
        width = max_h / aspect
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def set_cell(cell, text, bold=False, color=None, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_table(headers, rows, widths=None, center_cols=None, font=9.5):
    center_cols = center_cols or set()
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = True
    hdr = t.rows[0].cells
    for j, htext in enumerate(headers):
        set_cell(hdr[j], htext, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=font + 0.5,
                 align=WD_ALIGN_PARAGRAPH.CENTER if j in center_cols else None)
        shade(hdr[j], HEADER_FILL)
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            set_cell(cells[j], val, size=font,
                     align=WD_ALIGN_PARAGRAPH.CENTER if j in center_cols else None)
            if i % 2 == 1:
                shade(cells[j], ZEBRA)
    if widths:
        for row in t.rows:
            for j, wd in enumerate(widths):
                row.cells[j].width = Inches(wd)
    doc.add_paragraph()
    return t


def para(text=None, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    if text:
        p.add_run(text)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


# ============================ CONTENT ============================
t = doc.add_paragraph("Analytics & BI Platform Evaluation", style="Title")
sub = doc.add_paragraph()
r = sub.add_run("Prepared for Dollar Industries  |  Apparel & Textile Manufacturing and Distribution")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = TEAL
meta = doc.add_paragraph()
meta.add_run("Prepared by: [Consultant / Advisory]     ").font.size = Pt(9.5)
meta.add_run("Purpose: Evaluate three candidate platforms for enterprise data analytics and "
             "decision support, and recommend a fit-for-purpose path.").font.size = Pt(9.5)

# 1
doc.add_heading("1. Context & Evaluation Lens", level=1)
para("Dollar Industries operates a large manufacturing and multi-tier distribution business: "
     "factory output flows to depots, distributors, and dealers, and on to a wide retail and "
     "e-commerce network. The analytics platform must serve head-office leadership, regional and "
     "area sales managers, sales executives, and channel partners with reliable, fast insight.")
para("The platforms below are evaluated against the operational realities of this sector:")
add_table(
    ["Business Theme", "Representative Questions the Platform Must Answer"],
    [
        ["Primary sales (plant to distributor)", "Dispatch vs. plan by region, SKU, and period"],
        ["Secondary sales (distributor to retailer)", "Sell-through, slow movers, channel depth"],
        ["SKU & style hierarchy", "Performance by category / style / size / colour"],
        ["Distributor & dealer network", "Active vs. dormant partners, credit utilisation, outstanding receivables"],
        ["Trade schemes & incentives", "Scheme uptake, budget burn, ROI per scheme"],
        ["Inventory & stock movement", "Depot-wise stock, ageing, replenishment signals"],
        ["Sales targets", "Achievement by territory / ASM / sales executive"],
        ["Returns & quality complaints", "Defect trends, resolution turnaround"],
        ["Seasonality & demand", "Festive and seasonal demand planning"],
    ],
    widths=[2.3, 4.1],
)
para("The three candidates are evaluated on (1) Technical Feasibility & Future Scope, "
     "(2) Scalability, and (3) Costing.")

# 2
doc.add_heading("2. The Three Candidates at a Glance", level=1)
add_table(
    ["#", "Platform", "Category", "Interaction Model", "Ownership"],
    [
        ["A", "Tableau", "Traditional self-service BI", "Visual dashboards built by analysts", "Third-party licence (Salesforce)"],
        ["B", "AIBI (our platform)", "Conversational BI + Dashboards", "Natural-language chat and pinned KPI dashboards", "Our own IP - fully customisable"],
        ["C", "GenAI Conversational Analytics Assistant", "Agentic generative AI assistant", "NL chat with live charts, document Q&A, web lookup", "Custom build on a proven reference architecture"],
    ],
    widths=[0.3, 1.5, 1.4, 1.9, 1.6],
    center_cols={0},
)
p = doc.add_paragraph()
p.add_run("Note on Candidate C: ").bold = True
p.add_run("The reference documentation provided originates from a deployment in a different sector. "
          "In this evaluation it has been restructured and re-scoped to the apparel & textile context; "
          "the other client is not named. It represents the most advanced, fully agentic end of the spectrum.")

# 3
doc.add_heading("3. Reference Architectures", level=1)

doc.add_heading("3.1  Candidate A - Tableau (Traditional Self-Service BI)", level=2)
add_image(f"{DIAG}\\arch1.png", caption="Figure 1. Tableau - analyst-driven dashboarding architecture")
para("Mature, analyst-driven dashboarding. Insight is pre-built by skilled authors; business users "
     "consume. Natural-language is an add-on (Pulse), not the core.", bold_lead="Essence: ")

doc.add_heading("3.2  Candidate B - AIBI (Conversational BI Dashboard) - Our Platform", level=2)
add_image(f"{DIAG}\\arch2.png", caption="Figure 2. AIBI - hybrid conversational BI and dashboard architecture")
para("A hybrid - business users get both persistent KPI dashboards and a conversational interface, "
     "with LLM assistance to generate SQL and charts. Excel-to-database onboarding via LLM schema "
     "inference makes it fast to stand up. We own the roadmap, so it can be tailored to Dollar "
     "Industries' exact distributor, scheme, and SKU workflows.", bold_lead="Essence: ")

doc.add_heading("3.3  Candidate C - GenAI Conversational Analytics Assistant (Agentic)", level=2)
add_image(f"{DIAG}\\arch3.png", caption="Figure 3. GenAI Assistant - agentic, warehouse-backed architecture")
para("A fully agentic assistant. It reasons over a governed data warehouse (Bronze to Gold), answers "
     "from unstructured documents (catalogues, policies) via RAG, can perform web lookups, and streams "
     "live answers with charts. Highest capability ceiling; also the heaviest to build and run.",
     bold_lead="Essence: ")

# 4
doc.add_heading("4. Comparison - The Three Decision Criteria", level=1)

doc.add_heading("4.1  Technical Feasibility & Future Scope", level=2)
add_table(
    ["Dimension", "Tableau", "AIBI (ours)", "GenAI Assistant"],
    [
        ["Maturity / readiness", "Very high - proven, off-the-shelf", "Medium-high - productised core, evolving", "High - proven reference pattern, needs re-scoping"],
        ["Time to first value", "Fast for dashboards (with analysts)", "Fast - Excel/Postgres onboarding via LLM", "Moderate - requires ETL & data-model build"],
        ["Natural-language analytics", "Limited (Pulse add-on, Cloud-tied)", "Core capability - chat + dashboards", "Core capability - deep agentic chat"],
        ["Unstructured data (catalogues, policies)", "Not native", "Roadmap (extensible)", "Native via RAG"],
        ["Customisation to apparel workflows", "Within Tableau's framework only", "Full control - our IP", "Full control (bespoke build)"],
        ["Roadmap control", "Vendor-driven (Salesforce)", "We set the roadmap for the client", "Client / consultant-driven"],
        ["Future scope", "Incremental, vendor-paced", "Forecasting, connectors, agentic upgrades", "Voice, predictive ML, multimodal RAG, data-level access"],
        ["Key risk", "Conversational gap; ecosystem lock-in", "Product maturing; needs eng. investment", "Highest build complexity & integration effort"],
    ],
    widths=[1.7, 1.6, 1.6, 1.6], font=9,
)
para("Tableau is the safest for classic dashboards but weakest on conversation and customisation. "
     "AIBI offers the best balance of fast value, conversational UX, and full ownership. The GenAI "
     "Assistant has the highest ceiling but the steepest build.", bold_lead="Verdict: ")

doc.add_heading("4.2  Scalability", level=2)
add_table(
    ["Dimension", "Tableau", "AIBI (ours)", "GenAI Assistant"],
    [
        ["Proven user scale", "Enterprise-grade (thousands)", "Mid-market to enterprise (containerised)", "Cloud-native, auto-scaling (serverless)"],
        ["Data volume handling", "Strong (Hyper engine, extracts)", "Scales with PostgreSQL + tuning", "Highest - warehouse + parallel ETL"],
        ["Concurrency model", "Well understood; scale Server/Cloud", "App tier scales horizontally", "Compute scales independently"],
        ["Channel-partner reach (many viewers)", "Possible but licence-gated per viewer", "No per-seat lock-in", "No per-seat lock-in"],
        ["Real-time vs. batch", "Extract refresh / live connect", "Near-real-time queries (LLM latency)", "Streaming + daily incremental ETL"],
        ["Scaling constraint", "Cost grows linearly with users", "LLM latency/cost per query; needs caching", "LLM token throughput/quota; usage cost"],
    ],
    widths=[1.7, 1.6, 1.6, 1.6], font=9,
)
para("All three scale to Dollar Industries' size. Tableau scales but cost scales with every viewer - "
     "a real factor given a large distributor and field-sales base. AIBI and the GenAI Assistant avoid "
     "per-seat lock-in; the GenAI Assistant has the highest raw data-scale ceiling.", bold_lead="Verdict: ")

doc.add_heading("4.3  Costing", level=2)
para("Figures are indicative list prices for direction only; actuals depend on negotiation, FX, "
     "deployment model, and usage. INR equivalents assume approx. Rs. 84/USD.")
add_table(
    ["Cost Component", "Tableau", "AIBI (ours)", "GenAI Assistant"],
    [
        ["Licensing model", "Per-user subscription (recurring, USD)", "No 3rd-party licence - our IP", "No 3rd-party licence - custom build"],
        ["Indicative seat cost", "Creator ~$75, Explorer ~$42, Viewer ~$15 /user/mo (annual)", "Project + subscription to us (no per-seat tax)", "Project + run-cost (no per-seat tax)"],
        ["Infrastructure", "Tableau Server (on-prem) or Cloud", "PostgreSQL + app containers", "Containers + DB + object storage + vector search"],
        ["AI / LLM running cost", "None at core (Pulse extra)", "Per-token (low-cost small models) + caching", "Per-token + RAG indexing (scales with usage)"],
        ["Build / customisation", "Low (config) + analyst salaries", "Moderate (our platform, tailored)", "Highest (multi-agent, ETL, RAG)"],
        ["Cost behaviour as you grow", "Linear with user count", "Mostly fixed + variable LLM", "Variable, usage-driven"],
        ["Relative 3-year TCO", "Medium-High (seat-driven)", "Low-Medium (best value)", "High (highest capability)"],
    ],
    widths=[1.5, 1.7, 1.6, 1.7], font=9,
)
para("For an organisation with many light-touch users (field sales, distributors), Tableau's "
     "per-viewer licensing makes TCO climb quickly. AIBI offers the strongest value - no per-seat "
     "lock-in, controllable LLM cost via small models and caching. The GenAI Assistant carries the "
     "highest TCO, justified only when document Q&A, web search, and deep agentic analysis are "
     "must-haves.", bold_lead="Verdict: ")

# 5
doc.add_heading("5. Consolidated Scorecard", level=1)
para("Rated 1 (low) to 5 (high) for Dollar Industries' needs. Effort/cost rows: lower is better.")
star = lambda n: "★" * n + "☆" * (5 - n)
add_table(
    ["Criterion", "Tableau", "AIBI (ours)", "GenAI Assistant"],
    [
        ["Technical feasibility (today)", star(5), star(4), star(3)],
        ["Future scope / ceiling", star(3), star(4), star(5)],
        ["Conversational / NL analytics", star(2), star(4), star(5)],
        ["Customisation to your workflows", star(2), star(5), star(5)],
        ["Scalability (users)", star(5), star(4), star(5)],
        ["Scalability (data volume)", star(4), star(3), star(5)],
        ["Cost efficiency (TCO)", star(3), star(5), star(2)],
        ["Build / integration effort", "Low", "Medium", "High"],
        ["Vendor lock-in risk", "High", "Low", "Low"],
    ],
    widths=[2.2, 1.4, 1.4, 1.4], center_cols={1, 2, 3}, font=9.5,
)

# 6
doc.add_heading("6. Industry Use-Case Fit", level=1)
add_table(
    ["Use Case (Apparel & Textile)", "Best Served By"],
    [
        ["Standard executive dashboards (sales, stock, targets)", "Tableau or AIBI"],
        ["\"Ask a question\" self-serve for ASMs / sales executives", "AIBI / GenAI Assistant"],
        ["Distributor & dealer credit / outstanding monitoring", "AIBI (KPI + chat drill-down)"],
        ["Scheme ROI & budget-burn analysis", "AIBI / GenAI Assistant"],
        ["SKU / style / size / colour deep-dives on demand", "AIBI (chat + pinned KPIs)"],
        ["Querying product catalogues, policy & compliance docs", "GenAI Assistant (RAG)"],
        ["Blending internal data with market/web context", "GenAI Assistant (web search)"],
        ["Large read-only access for channel partners", "AIBI / GenAI (no per-seat licence)"],
    ],
    widths=[4.0, 2.4],
)

# 7
doc.add_heading("7. Recommendation - A Phased Path", level=1)
para("A single platform need not win outright. The pragmatic recommendation:")
add_image(f"{DIAG}\\roadmap.png", max_h=3.0, caption="Figure 4. Recommended phased adoption path")
numbered("Lowest TCO, no per-seat lock-in, conversational and dashboards on day one using existing "
         "PostgreSQL/Excel data. Best ratio of value to effort, fully tailorable to Dollar Industries.",
         bold_lead="Phase 1 - Land fast with AIBI. ")
numbered("Introduce a governed datamart (facts/dimensions for primary/secondary sales, schemes, stock) "
         "to improve answer quality and trust.", bold_lead="Phase 2 - Strengthen the data foundation. ")
numbered("Graduate to the agentic GenAI Assistant where document Q&A (catalogues, compliance), web "
         "context, and predictive analytics become priorities.", bold_lead="Phase 3 - Advanced. ")
numbered("For a small set of fixed, pixel-perfect boardroom reports, kept to a minimal licence footprint "
         "to contain cost.", bold_lead="Tableau - optional co-existence. ")
p = doc.add_paragraph()
p.add_run("Bottom line: ").bold = True
p.add_run("AIBI is the recommended primary platform for technical fit, scalability at this scale, and "
          "cost efficiency - with a clear, owned upgrade path into full agentic GenAI as needs mature.")

# 8
doc.add_heading("8. Assumptions & Notes", level=1)
bullet("Pricing is indicative list pricing for directional comparison only; final figures depend on "
       "commercial terms, FX, deployment (cloud vs. on-prem), and usage volumes.")
bullet("LLM running costs assume use of cost-efficient small models with caching and per-user rate limiting.")
bullet("Candidate C's capabilities are drawn from a restructured reference architecture originally "
       "deployed in another sector; no third-party client is named, and all examples are re-scoped to "
       "apparel & textile operations.")
bullet("Security baseline assumed across B and C: enterprise SSO, role-based access, read-only query "
       "execution, and least-privilege service identities.")

doc.save(OUT)
print("SAVED:", OUT)
